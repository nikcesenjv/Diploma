# Digitalizacija beležk SHS in Kraljevine Jugoslavije - Diplomsko delo
# Nik Česenj Vodovnik, 04180450 - Upravna informatika
# Študijsko leto 2022/2023
# Datoteka parlamint_objects_management.py

from docx.text.paragraph import Paragraph

from .parlamint_attendee_management import is_attendee

from src.logging import log

from src.management.json_management import parse_json
from src.management.text_management import *
from src.management.parlamint_objects_management import find_attendee
from src.management.path_management import parse_path

from src.objects.general_objects import File
from src.objects.parlamint_objects import ParlamintDocument, ParlamintAttendee, ParlamintHead, \
    ParlamintNote, ParlamintSpeaker, ParlamintSpeakerList, ParlamintUtterance

def create_parlamint_document(file: File, attendees: list[ParlamintAttendee]) -> ParlamintDocument:
    parlamint_document = ParlamintDocument(file)
    docx_paragraphs = parlamint_document.docx_document.paragraphs

    json_path = parse_path("path.parlamint_properties", f"parlamint_properties_{file.assembly}.json")
    properties = parse_json(json_path)["properties"]

    parlamint_document.elements, parlamint_document.attendees, all_attendees = \
        parse_parlamint_properties(properties, docx_paragraphs, attendees)

    parlamint_document.utterance_num, parlamint_document.segment_num = assign_document_statistics(parlamint_document)

    for attendee in parlamint_document.attendees:
        if not find_attendee(attendee.name, all_attendees):
            all_attendees.append(attendee)

    return parlamint_document, all_attendees

def assign_document_statistics(parlamint_document: ParlamintDocument) -> tuple[int, int]:
    utterance_num, segment_num = 1, 1

    for element_list in parlamint_document.elements:
        if not type(element_list) == ParlamintSpeakerList:
            continue

        for element in element_list.parlamint_speakers:
            if not type(element) == ParlamintSpeaker:
                continue

            element.utterance.document_id = parlamint_document.document_id
            element.utterance.utterance_num = utterance_num
            utterance_num += 1

            element.utterance.segment_start_num = segment_num
            segment_num += len(element.utterance.paragraphs)

    return utterance_num - 1, segment_num

def parse_parlamint_properties(properties: dict[str, str | list], paragraphs: list[Paragraph],
                               attendees: list[ParlamintAttendee]) \
        -> tuple[list[ParlamintHead | ParlamintNote], list[ParlamintAttendee], list[ParlamintAttendee]]:

    element_list, parlamint_attendee_list, parlamint_element, parlamint_speakers = [], [], None, False
    paragraphs = [parse_string(paragraph.text) for paragraph in paragraphs if len(paragraph.text) > 0]

    for i, property_element in enumerate(properties):
        match property_element["tag"]:
            case "head":
                parlamint_element = ParlamintHead(property_element["type"])
            case "note":
                parlamint_element = ParlamintNote(property_element["type"])
            case "speakers":
                parlamint_element = ParlamintSpeakerList()

        try:
            property_elements = properties[i + 1], properties[i + 2]
            parlamint_element.paragraphs, paragraphs = parse_parlamint_element(property_elements, paragraphs)

            if "role" in property_element:
                parlamint_attendee_list.extend(parse_parlamint_attendees(property_element, attendees,
                                                                         parlamint_element.paragraphs))
            if len(parlamint_element.paragraphs) > 0:
                element_list.append(parlamint_element)

        except IndexError:
            if not parlamint_speakers:
                parlamint_element.parlamint_speakers, paragraphs = parse_parlamint_speakers(paragraphs,
                                                                                            parlamint_attendee_list)
                element_list.append(parlamint_element)
                parlamint_speakers = True
            else:
                parlamint_element.paragraphs, paragraphs = parse_parlamint_element(property_element, paragraphs)
                element_list.append(parlamint_element)
    return element_list, parlamint_attendee_list, attendees

def parse_parlamint_element(properties: tuple[dict[str, str | list], dict[str, str | list]],
                            paragraphs: list[str]) -> list[str]:

    element_paragraphs, p = [], []
    paragraphs = [parse_string(paragraph) for paragraph in paragraphs if len(paragraph) > 0]

    for i, paragraph in enumerate(paragraphs):
        try:
            if closest_substring(paragraph, " ".join(properties[1]["similar words"])) is not None:
                return element_paragraphs, paragraphs[i:]
            elif closest_substring(paragraph, " ".join(properties[0]["similar words"])) is None:
                element_paragraphs.append(paragraph)
            else:
                return element_paragraphs, paragraphs[i:]

        except KeyError:
            if is_close_match_list(paragraph, properties["similar words"]):
                element_paragraphs.append(paragraph)
            elif closest_substring(paragraph, " ".join(properties[0]["similar words"])) is None:
                element_paragraphs.append(paragraph)

            return element_paragraphs, []

def parse_parlamint_attendees(property_element: dict[str, str | list], attendees: list[ParlamintAttendee],
                              paragraphs: list[Paragraph]) -> list[ParlamintAttendee]:

    parlamint_document_attendees, attendee_type = [], None
    attendee_names = remove_close_matches(" ".join(paragraphs), property_element["similar words"])

    for attendee_name in attendee_names.replace("\n", " ").split(", "):
        if is_attendee(attendee_name, attendees):
            continue

        if closest_substring(attendee_name, "izvestilac") is not None:
            attendee_name, attendee_type = remove_close_matches(attendee_name, ["izvestilac"]), "reporter"
        else:
            attendee_type = property_element["role"]

        parlamint_document_attendees.append(create_parlamint_attendee(attendee_name, attendee_type))

    return parlamint_document_attendees

def create_parlamint_attendee(name: str, attendee_type: str):
    return ParlamintAttendee(name, attendee_type)

def create_parlamint_speaker(speaker: str, parlamint_utterance: ParlamintUtterance):
    parlamint_note = create_parlamint_note("speaker", speaker)
    return ParlamintSpeaker(parlamint_note, parlamint_utterance)

def create_parlamint_note(note_type: str, string: str):
    parlamint_note = ParlamintNote(note_type)
    parlamint_note.add_paragraph(string)
    return parlamint_note

def create_parlamint_utterance(attendee: ParlamintAttendee):
    parlamint_utterance = ParlamintUtterance(attendee)
    return parlamint_utterance

def add_to_speakers_list(speakers_list, current_speaker, current_reader):
    speakers_list.append(current_speaker) if current_speaker is not None else None
    speakers_list.append(current_reader) if current_reader is not None else None
    return speakers_list

def add_to_speaker_reader(current_speaker, current_reader, string):
    current_speaker.utterance.add_paragraph(string) if current_speaker is not None else None
    current_reader.add_paragraph(string) if current_reader is not None else None
    return current_speaker, current_reader

def parse_parlamint_speakers(paragraphs: list[Paragraph], attendees: list[ParlamintAttendee]) -> list[ParlamintSpeaker]:
    speakers_list, current_speaker, current_reader = [], None, None

    for i, paragraph in enumerate([paragraph for paragraph in paragraphs[:-1] if paragraph != ""]):
        print(i)
        paragraph = parse_string(paragraph)

        if closest_substring(paragraph, "čita") is not None:
            speakers_list = add_to_speakers_list(speakers_list, current_speaker, current_reader)
            current_speaker, current_reader = None, None

            current_reader = create_parlamint_note("reading", paragraph.replace(":", "", 1))

        elif ":" in paragraph and ":" != paragraph[-1]:
            speaker, text = paragraph.split(":", 1)
            attendee = find_attendee(speaker, attendees)

            if attendee is not None:
                speakers_list = add_to_speakers_list(speakers_list, current_speaker, current_reader)
                current_speaker, current_reader = None, None

                parlamint_utterance = create_parlamint_utterance(attendee)
                parlamint_utterance.add_paragraph(text)

                current_speaker = create_parlamint_speaker(speaker, parlamint_utterance)
                continue

            current_speaker, current_reader = add_to_speaker_reader(current_speaker, current_reader, paragraph)

        else:
            if current_speaker is not None:
                current_speaker.utterance.add_paragraph(paragraph)
            else:
                current_reader.add_paragraph(paragraph)

    speakers_list = add_to_speakers_list(speakers_list, current_speaker, current_reader)
    return speakers_list, [paragraphs[-1]]
