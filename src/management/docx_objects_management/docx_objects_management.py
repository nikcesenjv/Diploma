# Digitalizacija beležk SHS in Kraljevine Jugoslavije - Diplomsko delo
# Nik Česenj Vodovnik, 04180450 - Upravna informatika
# Študijsko leto 2022/2023
# Datoteka docx_objects_management.py

from .docx_document_management import *
from .docx_row_management import *

from src.management.json_management import retrieve_json_value
from src.management.path_management import parse_path

from docx import Document

from src.objects.docx_objects import DocxDocument, DocxParagraph, DocxRow
from src.objects.general_objects import Document

def retrieve_namespace() -> dict[str, str]:
    return retrieve_json_value(parse_path("full_path.project", "path.docx_properties"), "NAMESPACE")

def create_docx_paragrahs(document_paragraph: str):
    ...

def create_docx_object(file: Document):
    docx_document = DocxDocument(file)

    document = Document(parse_path("full_path.project", "lib/documents", file.word_path))

    # docx_document = DocxDocument(file)
    # print(docx_document.to_string())

    # docx_document = Document(parse_path("full_path.project", file.word_path))

    """docx_document, namespace = DocxDocument(file), retrieve_namespace()

    for docx_paragraph in document_paragraphs(file, namespace):
        docx_document.add_paragraphs(create_docx_paragraphs(docx_paragraph, namespace))

    docx_document.text = docx_document.to_string()"""
    # print(docx_document.text)

    """tree = ET.ElementTree(docx_document.print_xml())
    ET.indent(tree, space="    ")
    tree.write("/Users/nikcesenjvodovnik/Documents/Programiranje/Diploma/test4.xml", encoding="UTF-8", xml_declaration=True)"""

    # return docx_document
    # return Document(parse_path("full_path.project", "lib/documents", file.word_path))

"""def create_docx_paragraphs(paragraph: ET, namespace: dict[str, str]) -> list[DocxParagraph]:
    rows = [create_docx_row(row, namespace) for row in paragraph.findall(".//w:r", namespace)]

    docx_paragraphs, same_font_list = [], []
    for row in rows:
        if len(same_font_list) == 0 or same_font_list[-1].font_size == row.font_size:
            same_font_list.append(row)
        else:
            docx_paragraphs.append(create_docx_paragraph(paragraph, same_font_list))
            same_font_list = [row]

    docx_paragraphs.append(create_docx_paragraph(paragraph, same_font_list))

    return [docx_paragraph for docx_paragraph in docx_paragraphs if docx_paragraph.text != ""]

def create_docx_paragraph(paragraph: ET, rows: list[DocxRow]) -> DocxParagraph:
    docx_paragraph = DocxParagraph(paragraph)
    docx_paragraph.rows = rows
    docx_paragraph.text = docx_paragraph.to_string()
    return docx_paragraph

def split_docx_paragraph(docx_paragraph: DocxParagraph, split_position: int):
    ...

def create_docx_row(row: ET, namespace: dict[str, str]) -> DocxRow:
    docx_row = DocxRow(row)
    docx_row.bold, docx_row.style, docx_row.font_size = parse_docx_row_properties(row, namespace)
    docx_row.text = docx_row.to_string()
    return docx_row"""
