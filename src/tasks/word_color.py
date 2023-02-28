from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement

import zipfile

doc = Document("/Users/nikcesenjvodovnik/Documents/Programiranje/Diploma/lib/documents/word/19_1932_SKJ/18-40_redni/7_XXIV_SKJ_redni_14.4.1932.docx")

styles = doc.styles

s = [s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]

for style in s:
    print(style)

x = Document()
paragraph_r = x.add_paragraph("To je test")

for p in doc.paragraphs:
    print(p.style)
    print(p.text)

# x.save("test.docx")
