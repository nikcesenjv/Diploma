# Digitalizacija beležk SHS in Kraljevine Jugoslavije - Diplomsko delo
# Nik Česenj Vodovnik, 04180450 - Upravna informatika
# Študijsko leto 2022/2023
# Datoteka docx_document.py

import xml.etree.ElementTree as ET

from zipfile import ZipFile
from docx import Document

from .docx_paragraph import DocxParagraph

from src.objects.general_objects import File

from src.management.path_management import parse_path

class DocxDocument:

    root_path = "/Users/nikcesenjvodovnik/Documents/Programiranje/Diploma/lib/documents/"

    def __init__(self, file: File):
        self._file: File = file

        self._paragraphs: list[DocxParagraph] = []
        self._text: str = None

    # GETTERS & SETTERS
    @property
    def file(self) -> File:
        return self._file

    @file.setter
    def file(self, new_file: File) -> None:
        self._file = new_file

    @property
    def paragraphs(self) -> list[DocxParagraph]:
        return self._paragraphs

    @paragraphs.setter
    def paragraphs(self, new_paragraphs: list[DocxParagraph]) -> None:
        self._paragraphs = new_paragraphs

    def add_paragraph(self, new_paragraph: DocxParagraph) -> None:
        self._paragraphs.append(new_paragraph)

    def add_paragraphs(self, new_paragraphs: list[DocxParagraph]) -> None:
        self._paragraphs.extend(new_paragraphs)

    """def to_string(self):
        return "\n".join([paragraph.text for paragraph in self.document.paragraphs])"""

    """

    @property
    def text(self) -> str:
        return self._text

    @text.setter
    def text(self, new_text: str) -> None:
        self._text = new_text

    def parse_paragraphs(self):
        return [DocxParagraph(paragraph) for paragraph in self.document.paragraphs]"""

"""    def to_string(self) -> str:
        return " ".join([paragraph.text for paragraph in self.paragraphs])

    def print_xml(self) -> ET:
        document = ET.fromstring(ZipFile(self.root_path + self.file.word_path).read("word/document.xml"))
        return document"""
"""from xml.dom import minidom
x = minidom.parseString(ET.tostring(document))
return x.toprettyxml()"""
