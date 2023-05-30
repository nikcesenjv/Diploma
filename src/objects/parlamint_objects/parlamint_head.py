# Digitalizacija beležk SHS in Kraljevine Jugoslavije - Diplomsko delo
# Nik Česenj Vodovnik, 04180450 - Upravna informatika
# Študijsko leto 2022/2023
# Datoteka parlamint_head.py

import xml.etree.ElementTree as ET

from docx.text.paragraph import Paragraph

from src.management.text_management import cyrillic_to_latin_text

class ParlamintHead:
    def __init__(self, head_type: str):
        self._head_type: str = head_type

        self._paragraphs: list[Paragraph] = []
        self._text: str = None

    @property
    def head_type(self) -> str:
        return self._head_type

    @head_type.setter
    def head_type(self, new_head_type: str) -> None:
        self._head_type = new_head_type

    @property
    def paragraphs(self) -> list[Paragraph]:
        return self._paragraphs

    @paragraphs.setter
    def paragraphs(self, new_paragraphs: list[Paragraph]) -> None:
        self._paragraphs = new_paragraphs

    def add_paragraph(self, new_paragraph: Paragraph) -> None:
        self._paragraphs.append(new_paragraph)

    def add_paragraphs(self, new_paragraphs: list[Paragraph]) -> None:
        self._paragraphs.extend(new_paragraphs)

    @property
    def text(self) -> str:
        return self._text

    @text.setter
    def text(self, new_text: str) -> None:
        self._text = new_text

    def to_string(self) -> str:
        self.text = " ".join([paragraph.replace("\n", " ") for paragraph in self.paragraphs])
        return self.text

    def to_element(self, parent_element: ET):
        head = ET.SubElement(parent_element, "head")

        if self.head_type is not None:
            head.set("type", self.head_type)

        head.text = self.to_string()
